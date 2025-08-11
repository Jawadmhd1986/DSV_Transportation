from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
import io, os, re, csv

app = Flask(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Money helpers
# ──────────────────────────────────────────────────────────────────────────────
DEC = Decimal

def q2d(x, default="0"):
    try:
        return DEC(str(x))
    except Exception:
        return DEC(default)

def money(d):
    if d is None:
        return ""
    if not isinstance(d, Decimal):
        d = q2d(d, "0")
    return f"{d.quantize(DEC('0.01'), rounding=ROUND_HALF_UP):,.2f}"

# ──────────────────────────────────────────────────────────────────────────────
# Normalization
# ──────────────────────────────────────────────────────────────────────────────
def norm_city(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.replace("_", " ").replace("–", "-").replace("—", "-")
    return s

def norm_truck(s: str) -> str:
    s = (s or "").strip().lower()
    aliases = {
        "curtainside": "box",
        "box / curtainside (5–10 tons)": "box",
        "box/curtainside": "box",
        "curtain side trailer": "box",
        "3 ton pickup": "city", "3 ton": "city", "1–3 ton": "city", "1-3 ton": "city",
        "7 ton pickup": "10_ton", "10 ton": "10_ton", "10-ton truck": "10_ton",
        "low bed": "lowbed", "low-bed": "lowbed",
        "double trailer": "double_trailer",
        "reefer": "reefer", "refrigerated": "reefer",
        "flat bed": "flatbed", "flat-bed": "flatbed",
        "tipper": "tipper", "dump": "tipper",
    }
    return aliases.get(s, s)

# ──────────────────────────────────────────────────────────────────────────────
# Robust rates loader
# Accepts:
#   A) origin,destination,truck_type,general_rate,chemical_rate
#   B) Delivery Locations,Vehicle Type,Price  (maps Hazmat FB -> chemical)
# Detects delimiters: comma / semicolon / tab
# ──────────────────────────────────────────────────────────────────────────────
def _detect_delimiter(sample: str) -> str:
    candidates = [",", ";", "\t", "|"]
    best = ","
    best_cols = 1
    first_line = sample.splitlines()[0] if sample else ""
    for d in candidates:
        cols = [c for c in first_line.split(d) if c.strip() != ""]
        if len(cols) > best_cols:
            best_cols = len(cols)
            best = d
    return best

def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
        txt = f.read()
    # If each line is quoted like "a,b,c", strip outer quotes
    lines = txt.splitlines()
    if lines and all(len(ln) > 2 and ln[0] == '"' and ln[-1] == '"' for ln in lines[: min(10, len(lines))]):
        txt = "\n".join(ln[1:-1] for ln in lines)
    return txt

def load_rates():
    rates = {}
    csv_path = os.path.join(app.root_path, "transport_rates.csv")
    if not os.path.exists(csv_path):
        print("[transport] rates CSV not found:", csv_path)
        return rates

    txt = _read_text(csv_path)
    delim = _detect_delimiter(txt)
    reader = csv.DictReader(txt.splitlines(), delimiter=delim)

    # Normalize headers
    header_map = {}
    for k in (reader.fieldnames or []):
        lk = (k or "").strip().lower()
        if lk in ("origin",): header_map[k] = "origin"
        elif lk in ("destination", "delivery locations", "delivery location", "city", "to"): header_map[k] = "destination"
        elif lk in ("truck_type", "vehicle type", "truck type"): header_map[k] = "truck_type"
        elif lk in ("general_rate", "general", "price", "rate"): header_map[k] = "general_rate"
        elif lk in ("chemical_rate", "chemical", "hazmat", "hazmat rate"): header_map[k] = "chemical_rate"
        else: header_map[k] = lk

    for row in reader:
        r = { header_map.get(k,k): (row.get(k,"") if row.get(k) is not None else "") for k in row }

        origin = norm_city(r.get("origin") or "*")
        dest   = norm_city(r.get("destination"))
        t_raw  = r.get("truck_type") or ""
        ttype  = norm_truck(t_raw)

        gen_raw = (r.get("general_rate") or "").strip()
        haz_raw = (r.get("chemical_rate") or "").strip()

        # 3‑column list handling
        if not gen_raw and not haz_raw and t_raw.strip().lower() == "hazmat fb" and r.get("general_rate") is not None:
            haz_raw = (r.get("general_rate") or "").strip()
        if not gen_raw and not haz_raw and r.get("general_rate") is not None and t_raw.strip().lower() != "hazmat fb":
            gen_raw = (r.get("general_rate") or "").strip()

        if not dest or not ttype:
            continue

        key = (origin or "*", dest)
        cell = rates.setdefault(key, {}).setdefault(ttype, {"general": None, "chemical": None})
        if gen_raw:
            cell["general"] = q2d(gen_raw)
        if haz_raw:
            cell["chemical"] = q2d(haz_raw)

    print(f"[transport] rates keys loaded: {len(rates)}")
    return rates

RATES = load_rates()

def lookup_rate(origin, destination, truck_type, cargo_type):
    cargo_key = "chemical" if (cargo_type or "").lower() == "chemical" else "general"
    o = norm_city(origin); d = norm_city(destination); t = norm_truck(truck_type)

    # Priority: (origin,destination) → ('*',destination) → (origin,'*')
    for key in ((o, d), ("*", d), (o, "*")):
        if key in RATES:
            data = RATES[key].get(t)
            if data:
                val = data.get(cargo_key) or data.get("general")
                if val is not None:
                    return val
    return None

# ──────────────────────────────────────────────────────────────────────────────
# Word helpers
# ──────────────────────────────────────────────────────────────────────────────
def replace_in_paragraph(paragraph, mapping):
    if not paragraph.runs:
        return
    original = "".join(r.text for r in paragraph.runs)
    replaced = original
    for k, v in mapping.items():
        replaced = replaced.replace(k, v)
    if replaced != original:
        for r in paragraph.runs:
            r.text = ""
        paragraph.runs[0].text = replaced

def replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

def find_details_table(doc):
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        header = " | ".join(c.text.strip() for c in tbl.rows[0].cells)
        if "Item" in header and "Unit Rate" in header and "Amount" in header:
            return tbl
    return None

def clear_table_body(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

def add_row(table, item, unit_rate="", amount=""):
    row = table.add_row()
    cells = row.cells
    if len(cells) >= 1: cells[0].text = str(item)
    if len(cells) >= 2: cells[1].text = str(unit_rate)
    if len(cells) >= 3: cells[2].text = str(amount)
    return row

def emphasize_row(row, font_pt=14):
    """Bold + bigger font for emphasis (Grand Total)."""
    for i, cell in enumerate(row.cells):
        for p in cell.paragraphs:
            # align last column (amount) to the right for clarity
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i == (len(row.cells)-1) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_pt)

# ──────────────────────────────────────────────────────────────────────────────
# UI
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/")
def home():
    return render_template("transport_form.html")

# Quick debug: /rates_debug?from=Mussafah&to=Mafraq&truck=flatbed&cargo=general
@app.route("/rates_debug")
def rates_debug():
    o = request.args.get("from", "")
    d = request.args.get("to", "")
    t = request.args.get("truck", "")
    c = (request.args.get("cargo", "general") or "general").lower()
    val = lookup_rate(o, d, t, c)
    return jsonify({
        "query": {"from": o, "to": d, "truck": t, "cargo": c},
        "normalized": {"from": norm_city(o), "to": norm_city(d), "truck": norm_truck(t)},
        "result": (float(val) if val is not None else None)
    })

# ──────────────────────────────────────────────────────────────────────────────
# Generate Transportation Quotation (NO VAT in table or totals)
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/generate_transport", methods=["POST"])
def generate_transport():
    origin         = (request.form.get("origin") or "").strip()
    destination    = (request.form.get("destination") or "").strip()
    trip_type      = (request.form.get("trip_type") or "one_way").strip()
    cargo_type     = (request.form.get("cargo_type") or "general").strip().lower()
    cicpa          = (request.form.get("cicpa") or "No").strip()
    stops          = [s.strip() for s in request.form.getlist("additional_cities[]") if s.strip()]
    truck_types    = request.form.getlist("truck_type[]") or []
    truck_qty_list = request.form.getlist("truck_qty[]") or []

    truck_labels = {
        "flatbed": "Flatbed (22–25T)",
        "box": "Box / Curtainside (5–10T)",
        "reefer": "Refrigerated (3–12T)",
        "city": "City (1–3T)",
        "tipper": "Tipper / Dump (15–20T)",
        "double_trailer": "Double Trailer",
        "10_ton": "10-Ton Truck",
        "lowbed": "Lowbed",
    }
    trip_label = {"one_way": "One Way", "back_load": "Back Load", "multi": "Multiple"}.get(trip_type, "One Way")

    # Build legs
    waypoints = [origin] + stops + [destination]
    legs = [(waypoints[i], waypoints[i+1]) for i in range(len(waypoints)-1) if waypoints[i] and waypoints[i+1]]
    if not legs:
        legs = [(origin, destination)]

    backload_mult = DEC("1.60") if trip_type == "back_load" else DEC("1.00")

    # Normalize chosen trucks
    chosen_trucks = []
    for t, q in zip(truck_types, truck_qty_list):
        nt = norm_truck(t)
        try:
            qty = int(float(q or "0"))
        except Exception:
            qty = 0
        if nt and qty > 0:
            chosen_trucks.append((nt, qty))

    subtotal = DEC("0")
    per_truck_rows = []

    for t, qty in chosen_trucks:
        label = truck_labels.get(t, t.title())
        combined = DEC("0")
        leg_descs = []

        for (frm, to) in legs:
            rate = lookup_rate(frm, to, t, cargo_type)
            if rate is None:
                continue  # show only priced legs
            leg_rate = (rate * backload_mult)
            combined += (leg_rate * qty)
            leg_descs.append(f"{frm} → {to}: AED {money(leg_rate)} x {qty}")

        if combined > 0 and leg_descs:
            unit_rate_str = f"AED {money((combined/qty) if qty else combined)}"
            amount_str    = f"AED {money(combined)}"
            per_truck_rows.append((f"{label} x {qty} — " + " | ".join(leg_descs), unit_rate_str, amount_str))
            subtotal += combined

    # Fees (NO VAT)
    total_trips = len(legs) * sum(q for _, q in chosen_trucks) if chosen_trucks else 0
    env_fixed   = DEC("10.00") * DEC(str(total_trips))    # AED 10 / trip / truck
    env_percent = subtotal * DEC("0.0015")                # 0.15% of invoice value
    grand_total = subtotal + env_fixed + env_percent      # ← VAT removed entirely

    # Summary placeholders used above the table
    truck_summary = "; ".join(f"{truck_labels.get(t, t.title())} x {qty}" for t, qty in chosen_trucks) or "N/A"
    route_str = " \u2192 ".join([p for p in [origin] + stops + [destination] if p]) or "N/A"
    general_flag  = "General Cargo" if cargo_type in ("general", "container") else ""
    chemical_flag = "Chemical Load" if cargo_type == "chemical" else ""

    tpl_path = os.path.join("templates", "TransportQuotation.docx")
    if not os.path.exists(tpl_path):
        return jsonify({"error": "TransportQuotation.docx not found under templates/"}), 500

    doc = Document(tpl_path)

    placeholders = {
        "{{TODAY_DATE}}": datetime.today().strftime("%d %b %Y"),
        "{{FROM}}":       origin or "N/A",
        "{{TO}}":         destination or "N/A",
        "{{TRUCK_TYPE}}": truck_summary,
        "{{GENERAL}}":    general_flag,
        "{{CHEMICAL}}":   chemical_flag,
        "{{TRIP_TYPE}}":  trip_label,
        "{{CICPA}}":      cicpa,
        "{{ROUTE}}":      route_str,
        "{{UNIT_RATE}}":  money(subtotal),
        "{{TOTAL_FEE}}":  money(grand_total)
    }
    replace_everywhere(doc, placeholders)

    # Rebuild the "Quotation Details" table:
    # 👉 ONLY priced truck lines + totals (NO VAT row)
    table = find_details_table(doc)
    if table:
        clear_table_body(table)

        # Add priced truck lines
        for desc, unit_rate, amount in per_truck_rows:
            add_row(table, desc, unit_rate, amount)

        # Totals (no VAT)
        add_row(table, "Subtotal (trips)", "", f"AED {money(subtotal)}")
        add_row(table, "Environmental Fee (AED 10 / trip / truck)", "", f"AED {money(env_fixed)}")
        add_row(table, "Environmental Levy (0.15% of invoice value)", "", f"AED {money(env_percent)}")

        # GRAND TOTAL emphasized
        gt_row = add_row(table, "GRAND TOTAL", "", f"AED {money(grand_total)}")
        emphasize_row(gt_row, font_pt=14)
    else:
        # Fallback if template table changes
        doc.add_paragraph("Quotation Details (Auto)")
        small = doc.add_table(rows=1, cols=3)
        hdr = small.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Item", "Unit Rate", "Amount (AED)"
        for desc, unit_rate, amount in per_truck_rows:
            add_row(small, desc, unit_rate, amount)
        add_row(small, "Subtotal (trips)", "", f"AED {money(subtotal)}")
        add_row(small, "Environmental Fee (AED 10 / trip / truck)", "", f"AED {money(env_fixed)}")
        add_row(small, "Environmental Levy (0.15% of invoice value)", "", f"AED {money(env_percent)}")
        gt_row = add_row(small, "GRAND TOTAL", "", f"AED {money(grand_total)}")
        emphasize_row(gt_row, font_pt=14)

    # Stream file
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    download_name = f"Transport_Quotation_{(origin or 'Origin').replace(' ','')}To{(destination or 'Destination').replace(' ','')}.docx"
    return send_file(buf, as_attachment=True, download_name=download_name)

# ──────────────────────────────────────────────────────────────────────────────
# Simple chat endpoint (frontend widget calls this)
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("message") or "").strip()
    if not text:
        return jsonify({"reply": "How can I help you with DSV transport, warehousing, or quotations?"})
    if re.match(r"^(hi|hello|hey)\b", text, re.I):
        return jsonify({"reply": "Hello! I’m here to help with storage or transport quotations. Ask me anything."})
    return jsonify({"reply": "Thanks for your message. For a quotation, fill the form and press Generate. If you need rates or specs, tell me which city/truck."})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
Generate a new quote and you’ll see:

no VAT row,

VAT not added to totals,

GRAND TOTAL rendered bold and larger.








You said:

ok keep the quotation table size alligned same with the detention table size, and the grand total make the font size 12

ChatGPT said:
Request for GPT-5 Pro
Reasoned for 17m 1s
done. this version makes the Quotation Details table the same overall width as the Detention Rates table (sets both to 100% page width), and sets the GRAND TOTAL row to bold 12 pt (amount right‑aligned). paste this full app.py:

python
Copy
Edit
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
import io, os, re, csv

app = Flask(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Money helpers
# ──────────────────────────────────────────────────────────────────────────────
DEC = Decimal

def q2d(x, default="0"):
    try:
        return DEC(str(x))
    except Exception:
        return DEC(default)

def money(d):
    if d is None:
        return ""
    if not isinstance(d, Decimal):
        d = q2d(d, "0")
    return f"{d.quantize(DEC('0.01'), rounding=ROUND_HALF_UP):,.2f}"

# ──────────────────────────────────────────────────────────────────────────────
# Normalization
# ──────────────────────────────────────────────────────────────────────────────
def norm_city(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.replace("_", " ").replace("–", "-").replace("—", "-")
    return s

def norm_truck(s: str) -> str:
    s = (s or "").strip().lower()
    aliases = {
        "curtainside": "box",
        "box / curtainside (5–10 tons)": "box",
        "box/curtainside": "box",
        "curtain side trailer": "box",
        "3 ton pickup": "city", "3 ton": "city", "1–3 ton": "city", "1-3 ton": "city",
        "7 ton pickup": "10_ton", "10 ton": "10_ton", "10-ton truck": "10_ton",
        "low bed": "lowbed", "low-bed": "lowbed",
        "double trailer": "double_trailer",
        "reefer": "reefer", "refrigerated": "reefer",
        "flat bed": "flatbed", "flat-bed": "flatbed",
        "tipper": "tipper", "dump": "tipper",
    }
    return aliases.get(s, s)

# ──────────────────────────────────────────────────────────────────────────────
# Rates loader (supports your 3‑column list & 5‑column matrix)
# ──────────────────────────────────────────────────────────────────────────────
def _detect_delimiter(sample: str) -> str:
    candidates = [",", ";", "\t", "|"]
    best = ","
    best_cols = 1
    first_line = sample.splitlines()[0] if sample else ""
    for d in candidates:
        cols = [c for c in first_line.split(d) if c.strip() != ""]
        if len(cols) > best_cols:
            best_cols = len(cols)
            best = d
    return best

def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8-sig", errors="ignore") as f:
        txt = f.read()
    lines = txt.splitlines()
    if lines and all(len(ln) > 2 and ln[0] == '"' and ln[-1] == '"' for ln in lines[: min(10, len(lines))]):
        txt = "\n".join(ln[1:-1] for ln in lines)
    return txt

def load_rates():
    rates = {}
    csv_path = os.path.join(app.root_path, "transport_rates.csv")
    if not os.path.exists(csv_path):
        print("[transport] rates CSV not found:", csv_path)
        return rates

    txt = _read_text(csv_path)
    delim = _detect_delimiter(txt)
    reader = csv.DictReader(txt.splitlines(), delimiter=delim)

    header_map = {}
    for k in (reader.fieldnames or []):
        lk = (k or "").strip().lower()
        if lk in ("origin",): header_map[k] = "origin"
        elif lk in ("destination", "delivery locations", "delivery location", "city", "to"): header_map[k] = "destination"
        elif lk in ("truck_type", "vehicle type", "truck type"): header_map[k] = "truck_type"
        elif lk in ("general_rate", "general", "price", "rate"): header_map[k] = "general_rate"
        elif lk in ("chemical_rate", "chemical", "hazmat", "hazmat rate"): header_map[k] = "chemical_rate"
        else: header_map[k] = lk

    for row in reader:
        r = { header_map.get(k,k): (row.get(k,"") if row.get(k) is not None else "") for k in row }

        origin = norm_city(r.get("origin") or "*")
        dest   = norm_city(r.get("destination"))
        t_raw  = r.get("truck_type") or ""
        ttype  = norm_truck(t_raw)

        gen_raw = (r.get("general_rate") or "").strip()
        haz_raw = (r.get("chemical_rate") or "").strip()

        # map 3‑column format (Price) to general or chemical
        if not gen_raw and not haz_raw and t_raw.strip().lower() == "hazmat fb" and r.get("general_rate") is not None:
            haz_raw = (r.get("general_rate") or "").strip()
        if not gen_raw and not haz_raw and r.get("general_rate") is not None and t_raw.strip().lower() != "hazmat fb":
            gen_raw = (r.get("general_rate") or "").strip()

        if not dest or not ttype:
            continue

        key = (origin or "*", dest)
        cell = rates.setdefault(key, {}).setdefault(ttype, {"general": None, "chemical": None})
        if gen_raw:
            cell["general"] = q2d(gen_raw)
        if haz_raw:
            cell["chemical"] = q2d(haz_raw)

    print(f"[transport] rates keys loaded: {len(rates)}")
    return rates

RATES = load_rates()

def lookup_rate(origin, destination, truck_type, cargo_type):
    cargo_key = "chemical" if (cargo_type or "").lower() == "chemical" else "general"
    o = norm_city(origin); d = norm_city(destination); t = norm_truck(truck_type)
    for key in ((o, d), ("*", d), (o, "*")):
        if key in RATES:
            data = RATES[key].get(t)
            if data:
                val = data.get(cargo_key) or data.get("general")
                if val is not None:
                    return val
    return None

# ──────────────────────────────────────────────────────────────────────────────
# Word helpers
# ──────────────────────────────────────────────────────────────────────────────
def replace_in_paragraph(paragraph, mapping):
    if not paragraph.runs:
        return
    original = "".join(r.text for r in paragraph.runs)
    replaced = original
    for k, v in mapping.items():
        replaced = replaced.replace(k, v)
    if replaced != original:
        for r in paragraph.runs:
            r.text = ""
        paragraph.runs[0].text = replaced

def replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

def find_details_table(doc):
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        header = " | ".join(c.text.strip() for c in tbl.rows[0].cells)
        if "Item" in header and "Unit Rate" in header and "Amount" in header:
            return tbl
    return None

def clear_table_body(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

def add_row(table, item, unit_rate="", amount=""):
    row = table.add_row()
    cells = row.cells
    if len(cells) >= 1: cells[0].text = str(item)
    if len(cells) >= 2: cells[1].text = str(unit_rate)
    if len(cells) >= 3: cells[2].text = str(amount)
    return row

def emphasize_row(row, font_pt=12):  # ← 12 pt as requested
    for i, cell in enumerate(row.cells):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i == (len(row.cells)-1) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_pt)

def set_table_width_pct(table, pct=100):
    """
    Force a table to occupy a percentage of the page width.
    Word uses fiftieths of a percent, so 100% = 5000.
    """
    w = str(int(pct * 50))  # 100% -> 5000
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tbl.tblPr = tblPr
    tblW = tblPr.tblW or OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), w)
    tblPr.append(tblW)

# ──────────────────────────────────────────────────────────────────────────────
# UI
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/")
def home():
    return render_template("transport_form.html")

@app.route("/rates_debug")
def rates_debug():
    o = request.args.get("from", "")
    d = request.args.get("to", "")
    t = request.args.get("truck", "")
    c = (request.args.get("cargo", "general") or "general").lower()
    val = lookup_rate(o, d, t, c)
    return jsonify({
        "query": {"from": o, "to": d, "truck": t, "cargo": c},
        "normalized": {"from": norm_city(o), "to": norm_city(d), "truck": norm_truck(t)},
        "result": (float(val) if val is not None else None)
    })

# ──────────────────────────────────────────────────────────────────────────────
# Generate Transportation Quotation (NO VAT; bold 12pt GRAND TOTAL; tables same width)
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/generate_transport", methods=["POST"])
def generate_transport():
    origin         = (request.form.get("origin") or "").strip()
    destination    = (request.form.get("destination") or "").strip()
    trip_type      = (request.form.get("trip_type") or "one_way").strip()
    cargo_type     = (request.form.get("cargo_type") or "general").strip().lower()
    cicpa          = (request.form.get("cicpa") or "No").strip()
    stops          = [s.strip() for s in request.form.getlist("additional_cities[]") if s.strip()]
    truck_types    = request.form.getlist("truck_type[]") or []
    truck_qty_list = request.form.getlist("truck_qty[]") or []

    truck_labels = {
        "flatbed": "Flatbed (22–25T)",
        "box": "Box / Curtainside (5–10T)",
        "reefer": "Refrigerated (3–12T)",
        "city": "City (1–3T)",
        "tipper": "Tipper / Dump (15–20T)",
        "double_trailer": "Double Trailer",
        "10_ton": "10-Ton Truck",
        "lowbed": "Lowbed",
    }
    trip_label = {"one_way": "One Way", "back_load": "Back Load", "multi": "Multiple"}.get(trip_type, "One Way")

    # Build legs
    waypoints = [origin] + stops + [destination]
    legs = [(waypoints[i], waypoints[i+1]) for i in range(len(waypoints)-1) if waypoints[i] and waypoints[i+1]]
    if not legs:
        legs = [(origin, destination)]

    backload_mult = DEC("1.60") if trip_type == "back_load" else DEC("1.00")

    # Normalize chosen trucks
    chosen_trucks = []
    for t, q in zip(truck_types, truck_qty_list):
        nt = norm_truck(t)
        try:
            qty = int(float(q or "0"))
        except Exception:
            qty = 0
        if nt and qty > 0:
            chosen_trucks.append((nt, qty))

    subtotal = DEC("0")
    per_truck_rows = []

    for t, qty in chosen_trucks:
        label = truck_labels.get(t, t.title())
        combined = DEC("0")
        leg_descs = []

        for (frm, to) in legs:
            rate = lookup_rate(frm, to, t, cargo_type)
            if rate is None:
                continue  # show only priced legs
            leg_rate = (rate * backload_mult)
            combined += (leg_rate * qty)
            leg_descs.append(f"{frm} → {to}: AED {money(leg_rate)} x {qty}")

        if combined > 0 and leg_descs:
            unit_rate_str = f"AED {money((combined/qty) if qty else combined)}"
            amount_str    = f"AED {money(combined)}"
            per_truck_rows.append((f"{label} x {qty} — " + " | ".join(leg_descs), unit_rate_str, amount_str))
            subtotal += combined

    # Fees (NO VAT)
    total_trips = len(legs) * sum(q for _, q in chosen_trucks) if chosen_trucks else 0
    env_fixed   = DEC("10.00") * DEC(str(total_trips))    # AED 10 / trip / truck
    env_percent = subtotal * DEC("0.0015")                # 0.15% of invoice value
    grand_total = subtotal + env_fixed + env_percent

    # Summaries above the table
    truck_summary = "; ".join(f"{truck_labels.get(t, t.title())} x {qty}" for t, qty in chosen_trucks) or "N/A"
    route_str = " \u2192 ".join([p for p in [origin] + stops + [destination] if p]) or "N/A"
    general_flag  = "General Cargo" if cargo_type in ("general", "container") else ""
    chemical_flag = "Chemical Load" if cargo_type == "chemical" else ""

    tpl_path = os.path.join("templates", "TransportQuotation.docx")
    if not os.path.exists(tpl_path):
        return jsonify({"error": "TransportQuotation.docx not found under templates/"}), 500

    doc = Document(tpl_path)

    placeholders = {
        "{{TODAY_DATE}}": datetime.today().strftime("%d %b %Y"),
        "{{FROM}}":       origin or "N/A",
        "{{TO}}":         destination or "N/A",
        "{{TRUCK_TYPE}}": truck_summary,
        "{{GENERAL}}":    general_flag,
        "{{CHEMICAL}}":   chemical_flag,
        "{{TRIP_TYPE}}":  trip_label,
        "{{CICPA}}":      cicpa,
        "{{ROUTE}}":      route_str,
        "{{UNIT_RATE}}":  money(subtotal),
        "{{TOTAL_FEE}}":  money(grand_total)
    }
    replace_everywhere(doc, placeholders)

    # Quotation table → only priced rows + totals
    table = find_details_table(doc)
    if table:
        clear_table_body(table)
        for desc, unit_rate, amount in per_truck_rows:
            add_row(table, desc, unit_rate, amount)
        add_row(table, "Subtotal (trips)", "", f"AED {money(subtotal)}")
        add_row(table, "Environmental Fee (AED 10 / trip / truck)", "", f"AED {money(env_fixed)}")
        add_row(table, "Environmental Levy (0.15% of invoice value)", "", f"AED {money(env_percent)}")
        gt_row = add_row(table, "GRAND TOTAL", "", f"AED {money(grand_total)}")
        emphasize_row(gt_row, font_pt=12)

        # 🔧 make this table full page width (align size with other tables)
        set_table_width_pct(table, 100)
    else:
        doc.add_paragraph("Quotation Details (Auto)")
        small = doc.add_table(rows=1, cols=3)
        hdr = small.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Item", "Unit Rate", "Amount (AED)"
        for desc, unit_rate, amount in per_truck_rows:
            add_row(small, desc, unit_rate, amount)
        add_row(small, "Subtotal (trips)", "", f"AED {money(subtotal)}")
        add_row(small, "Environmental Fee (AED 10 / trip / truck)", "", f"AED {money(env_fixed)}")
        add_row(small, "Environmental Levy (0.15% of invoice value)", "", f"AED {money(env_percent)}")
        gt_row = add_row(small, "GRAND TOTAL", "", f"AED {money(grand_total)}")
        emphasize_row(gt_row, font_pt=12)
        set_table_width_pct(small, 100)

    # Also ensure ALL tables (including Detention Rates) are 100% width for visual alignment
    for t in doc.tables:
        set_table_width_pct(t, 100)

    # Stream file
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    download_name = f"Transport_Quotation_{(origin or 'Origin').replace(' ','')}To{(destination or 'Destination').replace(' ','')}.docx"
    return send_file(buf, as_attachment=True, download_name=download_name)

# ──────────────────────────────────────────────────────────────────────────────
# Minimal chat endpoint
# ──────────────────────────────────────────────────────────────────────────────
@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True, silent=True) or {}
    text = (data.get("message") or "").strip()
    if not text:
        return jsonify({"reply": "How can I help you with DSV transport, warehousing, or quotations?"})
    if re.match(r"^(hi|hello|hey)\b", text, re.I):
        return jsonify({"reply": "Hello! I’m here to help with storage or transport quotations. Ask me anything."})
    return jsonify({"reply": "Thanks for your message. For a quotation, fill the form and press Generate. If you need rates or specs, tell me which city/truck."})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
